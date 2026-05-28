#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def read_markdown_file(file_path):
    """Lê um arquivo Markdown"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def convert_md_to_docx(md_file_path, docx_file_path):
    """Converte um arquivo Markdown para DOCX"""
    
    # Ler o conteúdo do markdown
    content = read_markdown_file(md_file_path)
    
    # Criar documento Word
    doc = Document()
    
    # Processar linha por linha
    lines = content.split('\n')
    current_list = False
    
    for line in lines:
        line = line.rstrip()
        
        # Títulos H1
        if line.startswith('# '):
            heading = line[2:].strip()
            p = doc.add_heading(heading, level=1)
            continue
        
        # Títulos H2
        if line.startswith('## '):
            heading = line[3:].strip()
            p = doc.add_heading(heading, level=2)
            continue
        
        # Títulos H3
        if line.startswith('### '):
            heading = line[4:].strip()
            p = doc.add_heading(heading, level=3)
            continue
        
        # Títulos H4
        if line.startswith('#### '):
            heading = line[5:].strip()
            p = doc.add_heading(heading, level=4)
            continue
        
        # Títulos H5
        if line.startswith('##### '):
            heading = line[6:].strip()
            p = doc.add_heading(heading, level=5)
            continue
        
        # Tabelas (simples)
        if line.startswith('|'):
            # Pular linhas de separação de tabela
            if re.match(r'^\|\s*-+', line):
                continue
            
            # Extrair células
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            
            # Criar tabela se não existir
            if not hasattr(doc, '_last_element') or not isinstance(doc._element[-1], type(None)):
                if not current_list and len(cells) > 0:
                    table = doc.add_table(rows=1, cols=len(cells))
                    table.style = 'Light Grid Accent 1'
                    current_list = True
                    
                    # Adicionar célula header
                    for i, cell_text in enumerate(cells):
                        table.rows[0].cells[i].text = cell_text
            
            continue
        
        # Linhas em branco
        if not line.strip():
            if not current_list:
                doc.add_paragraph()
            current_list = False
            continue
        
        # Parágrafos normais
        if line.strip():
            current_list = False
            
            # Remover caracteres de markdown simples
            text = line
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # **bold**
            text = re.sub(r'\*(.*?)\*', r'\1', text)      # *italic*
            text = re.sub(r'`(.*?)`', r'\1', text)        # `code`
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # [link](url)
            
            doc.add_paragraph(text)
    
    # Salvar documento
    doc.save(docx_file_path)
    print(f"✅ Convertido: {md_file_path} → {docx_file_path}")

# Diretórios
docs_dir = r'c:\Users\User\Downloads\ProjetoEstagio\docs'
output_dir = r'c:\Users\User\Downloads\ProjetoEstagio\docs'

# Converter arquivos
files_to_convert = [
    ('UC01_Gerenciar_Usuarios.md', 'UC01_Gerenciar_Usuarios.docx'),
    ('UC02_Simular_Pagamento.md', 'UC02_Simular_Pagamento.docx'),
]

print("Convertendo arquivos Markdown para DOCX...\n")

for md_file, docx_file in files_to_convert:
    md_path = os.path.join(docs_dir, md_file)
    docx_path = os.path.join(output_dir, docx_file)
    
    if os.path.exists(md_path):
        try:
            convert_md_to_docx(md_path, docx_path)
        except Exception as e:
            print(f"❌ Erro ao converter {md_file}: {str(e)}")
    else:
        print(f"❌ Arquivo não encontrado: {md_path}")

print("\n✨ Conversão concluída!")
print(f"📁 Arquivos salvos em: {output_dir}")
